import io
import os
import zipfile
from datetime import datetime

import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer

# ======================================================
# App Config
# ======================================================
st.set_page_config(page_title="SPC — PDF-like Merge + TOC (Titles) + Page Numbers", layout="wide")
st.title("📚 SPC — Gabung DOCX (PDF-like) + TOC Tajuk Paper + Muka Surat")
st.caption("TOC di atas fail, tajuk ambil daripada Heading 1 setiap paper. Setiap dokumen di halaman baharu. Format asal setiap fail dikekalkan.")

st.markdown("""
**Cara guna ringkas**
1) Sediakan **ZIP** yang mengandungi semua `.docx` dalam folder.  
2) Upload ZIP → klik **Gabungkan** → muat turun fail gabungan `.docx`.  
3) Buka di Microsoft Word → **Right-click** pada TOC → **Update Field** → **Update entire table** (untuk kemas kini nombor halaman & rupa laporan dengan dot leaders).

> **Nota penting**: TOC guna **tajuk paper** daripada *Heading 1* dalam setiap dokumen. Jika tiada Heading 1, sistem akan guna **nama fail** sebagai tajuk.
""")

# ======================================================
# Utilities
# ======================================================

def zip_docx_entries_in_order(zip_bytes: bytes):
    """
    Pulangkan senarai (name_in_zip, bytes) mengikut susunan entri dalam ZIP (ZipInfo order).
    Hanya .docx diambil; folder diabaikan.
    """
    result = []
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = info.filename
            if name.lower().endswith(".docx"):
                with zf.open(info, 'r') as fp:
                    result.append((name, fp.read()))
    return result  # kekalkan susunan asal dalam ZIP


def extract_title_from_doc_bytes(doc_bytes: bytes, fallback_name: str) -> str:
    """
    Cuba dapatkan tajuk daripada Heading 1 dalam dokumen.
    Jika tiada, guna nama fail (tanpa .docx) sebagai fallback.
    """
    try:
        d = Document(io.BytesIO(doc_bytes))
        for p in d.paragraphs:
            try:
                if p.style and p.style.name and str(p.style.name).lower().startswith("heading 1"):
                    txt = (p.text or "").strip()
                    if txt:
                        return txt
            except Exception:
                continue
    except Exception:
        pass
    # fallback: nama fail tanpa extension
    base = os.path.splitext(os.path.basename(fallback_name))[0]
    return base


def add_field_run_hidden(paragraph, instr_text: str):
    """
    Sisip *complex field* yang disembunyikan (hidden/vanish) dalam satu paragraph.
    Field code akan diset 'vanish' supaya tidak kelihatan pada paparan/print default.
    """
    # BEGIN
    r_begin = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    vanish = OxmlElement("w:vanish")
    rpr.append(vanish)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(rpr)
    r_begin.append(fld_begin)
    paragraph._p.append(r_begin)

    # INSTR
    r_instr = OxmlElement("w:r")
    rpr2 = OxmlElement("w:rPr")
    vanish2 = OxmlElement("w:vanish")
    rpr2.append(vanish2)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instr_text} "
    r_instr.append(rpr2)
    r_instr.append(instr)
    paragraph._p.append(r_instr)

    # END
    r_end = OxmlElement("w:r")
    rpr3 = OxmlElement("w:rPr")
    vanish3 = OxmlElement("w:vanish")
    rpr3.append(vanish3)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(rpr3)
    r_end.append(fld_end)
    paragraph._p.append(r_end)


def add_tc_hidden_entry(doc: Document, title: str, level: int = 1):
    """
    Tambah satu TC field tersembunyi (hidden) untuk TOC berdasarkan tajuk.
    Contoh instruksi:  TC "My Paper Title" \l 1
    """
    p = doc.add_paragraph()  # paragraph kosong khusus untuk TC
    safe_title = title.replace('"', "'")  # elak quote clash
    instr = f'TC "{safe_title}" \\l {level}'
    add_field_run_hidden(p, instr)


def add_toc_field_from_tc(doc: Document, title="Table of Contents"):
    """
    Sisip TOC yang membaca entri dari TC fields: { TOC \h \z \f "TC" }.
    Ini membolehkan kita kawal tajuk TOC ikut TC yang kita sisip (tajuk paper),
    tanpa perlu mengubah kandungan dokumen asal.
    """
    # Tajuk TOC
    t = doc.add_paragraph()
    run = t.add_run(title)
    run.bold = True
    # biar Word tentukan rupa; nombor kanan + dot leaders akan muncul selepas Update Field di Word
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Field TOC (berpunca dari TC fields)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    # \f "TC" => kumpul entri daripada TC fields
    fld.set(qn("w:instr"), 'TOC \\h \\z \\f "TC"')
    p._p.append(fld)


def add_page_numbers_all_sections(doc: Document):
    """
    Tambah 'Page X of Y' di footer (tengah) untuk semua seksyen dalam dokumen.
    Dilakukan selepas gabungan supaya merangkumi seksyen daripada sub-documents.
    """
    def add_field_run(paragraph, field):
        r_begin = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r_begin.append(fld_begin)
        paragraph._p.append(r_begin)

        r_instr = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field} "
        r_instr.append(instr)
        paragraph._p.append(r_instr)

        r_end = OxmlElement("w:r")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r_end.append(fld_end)
        paragraph._p.append(r_end)

    for section in doc.sections:
        footer = section.footer
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("Page ")
        add_field_run(para, "PAGE")
        para.add_run(" of ")
        add_field_run(para, "NUMPAGES")


def combine_pdf_like_with_toc_titles(zip_bytes: bytes) -> bytes:
    """
    Gabungkan dokumen .docx secara 'PDF-like' (setiap dokumen di halaman baharu),
    TOC di awal berdasarkan **tajuk kertas** (Heading 1) via TC fields tersembunyi,
    dan tambah muka surat untuk semua seksyen.
    """
    entries = zip_docx_entries_in_order(zip_bytes)
    if not entries:
        raise ValueError("ZIP tidak mengandungi sebarang .docx")

    # Dokumen asas
    base = Document()

    # 1) Kumpul tajuk dan sisip TC fields (hidden) untuk setiap dokumen
    titles = []
    for name, blob in entries:
        title = extract_title_from_doc_bytes(blob, name)
        titles.append(title)
        add_tc_hidden_entry(base, title, level=1)

    # 2) Sisip TOC di atas sekali, bersumberkan TC fields
    add_toc_field_from_tc(base, title="Table of Contents")

    # 3) Pisahkan TOC dari kandungan
    base.add_page_break()

    # 4) Gabungkan semua dokumen — setiap satu bermula di halaman baharu
    composer = Composer(base)
    for idx, (name, blob) in enumerate(entries):
        if idx > 0:
            base.add_page_break()  # pemisah halaman antara paper
        sub = Document(io.BytesIO(blob))
        composer.append(sub)

    # 5) Simpan sementara
    tmp = io.BytesIO()
    composer.save(tmp)
    tmp.seek(0)

    # 6) Tambah muka surat untuk semua seksyen pada fail komposit
    compiled = Document(tmp)
    add_page_numbers_all_sections(compiled)

    out = io.BytesIO()
    compiled.save(out)
    out.seek(0)
    return out.read()

# ======================================================
# UI
# ======================================================

st.subheader("Muat Naik ZIP Anda")
zip_file = st.file_uploader(
    "Upload satu ZIP (mengandungi folder + .docx). Susunan ikut folder (ZipInfo order).",
    type=["zip"],
    accept_multiple_files=False
)

st.info(
    "• TOC menggunakan tajuk kertas daripada Heading 1 bagi setiap dokumen (fallback: nama fail). "
    "Rupa laporan (tajuk kiri, nombor kanan, dot leaders) akan keluar selepas buka di Word dan Right-click TOC → Update Field → Update entire table.\n"
    "• Setiap dokumen bermula di halaman baharu (tak bercampur). Kandungan asal tidak diubah."
)

default_name = f"SPC_Proceedings_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
out_name = st.text_input("Nama fail output", value=default_name)

if st.button("🚀 Gabungkan (TOC tajuk paper, setiap dokumen halaman baharu)"):
    try:
        if not zip_file:
            st.warning("Sila upload satu fail ZIP.")
        else:
            with st.spinner("Menggabungkan dokumen ikut susunan folder dalam ZIP..."):
                combined_bytes = combine_pdf_like_with_toc_titles(zip_file.read())
            st.success("Siap! Muat turun di bawah. (Di Microsoft Word: Right-click TOC → Update Field → Update entire table)")
            st.download_button(
                "⬇️ Muat Turun Fail Gabungan",
                data=combined_bytes,
                file_name=out_name or "SPC_Proceedings.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    except Exception as e:
        st.error("Ralat semasa menggabungkan dokumen.")
        st.exception(e)
